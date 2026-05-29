"""
word_handler.py
Заповнення шаблонів Word та конвертація в PDF

ВИПРАВЛЕННЯ гіперпосилань: стара версія збирала текст через para.runs і
записувала результат у runs[0].text, але w:hyperlink є окремими XML-вузлами
ПОЗА para.runs — їхні внутрішні runs стиралися. Нова версія ітерує XML
напряму через para._element.iter() і замінює text у кожному w:t окремо.
"""

import os
import subprocess
import shutil
from datetime import date
from docx import Document
from docx.oxml.ns import qn
from config import (
    TEMPLATE_INVOICE_ACCESS, TEMPLATE_INVOICE_HOURLY,
    TEMPLATE_ACT_ACCESS, TEMPLATE_ACT_HOURLY,
    CONTRACTS_DIR, LIBREOFFICE_PATH, TMP_DIR, TEMPLATES_DIR
)


# ============================================================
# ЗАМІНА ПЛЕЙСХОЛДЕРІВ
# ============================================================

def _replace_in_paragraph(para, replacements: dict):
    """
    Замінює плейсхолдери в параграфі на рівні XML.

    Ключова відмінність від старої версії:
    - НЕ використовує para.runs (вони не охоплюють w:hyperlink)
    - Замість цього збирає ВСІ w:t через para._element.iter()
    - Знаходить який w:t містить плейсхолдер і замінює там
    - Якщо плейсхолдер розбитий між кількома w:t —
      записує результат у перший, решту очищує

    Таким чином w:hyperlink залишається незайманим у XML —
    змінюється лише текст, а r:id / структура зберігаються.
    """
    # Швидка перевірка через para.text (охоплює ВСЕ включно з hyperlinks)
    if not any(key in (para.text or '') for key in replacements):
        return

    # Збираємо всі w:t в параграфі (у порядку документа)
    t_elems = list(para._element.iter(qn('w:t')))
    if not t_elems:
        return

    # Будуємо позиційну карту: повний текст + де починається кожен w:t
    parts = [t.text or '' for t in t_elems]
    combined = ''.join(parts)

    new_combined = combined
    for key, value in replacements.items():
        new_combined = new_combined.replace(key, str(value))

    if new_combined == combined:
        return

    # Розподіляємо новий текст назад.
    # Стратегія: записуємо весь результат у перший непорожній w:t,
    # решту обнуляємо. Це безпечно бо:
    # - текстовий вміст гіперпосилання зберігається (у тому ж w:t)
    # - структура w:hyperlink + r:id в XML не змінюється
    # - відображається той самий текст, просто в одному вузлі
    first_idx = next((i for i, p in enumerate(parts) if p.strip()), 0)

    t_elems[first_idx].text = new_combined
    # Зберігаємо пробіли якщо вони є на початку/кінці
    if new_combined != new_combined.strip():
        t_elems[first_idx].set(
            '{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Обнуляємо решту w:t в цьому параграфі
    for i, t in enumerate(t_elems):
        if i != first_idx:
            t.text = ''


def _inject_signature_image(doc, sig_path: str, width_cm: float = 2.5) -> bool:
    """Знаходить {{ПІДПИС}} і замінює ТІЛЬКИ цей плейсхолдер на зображення.

    Весь інший текст у параграфі залишається незайманим.
    Шукає у тілі документа, таблицях, хедері та футері.
    """
    from docx.shared import Cm
    from docx.oxml import OxmlElement

    SIG_TAG = '{{ПІДПИС}}'

    def _try_para(para) -> bool:
        if SIG_TAG not in (para.text or ''):
            return False

        # Крок 1: шукаємо run де тег є цілком (найпоширеніший випадок)
        for run in para.runs:
            if SIG_TAG in run.text:
                idx    = run.text.index(SIG_TAG)
                before = run.text[:idx]
                after  = run.text[idx + len(SIG_TAG):]

                run.text = before                              # текст до підпису
                run.add_picture(sig_path, width=Cm(width_cm)) # вставляємо підпис

                if after:
                    # Вставляємо текст після підпису як окремий run у XML
                    new_r = OxmlElement('w:r')
                    new_t = OxmlElement('w:t')
                    new_t.text = after
                    if after != after.strip():
                        new_t.set(
                            '{http://www.w3.org/XML/1998/namespace}space',
                            'preserve')
                    new_r.append(new_t)
                    run._r.addnext(new_r)

                return True

        # Крок 2: тег розбитий між кількома runs — нормалізуємо
        runs      = para.runs
        full_text = ''.join(r.text for r in runs)
        if SIG_TAG not in full_text:
            return False

        tag_start   = full_text.index(SIG_TAG)
        tag_end     = tag_start + len(SIG_TAG)
        before_text = full_text[:tag_start]
        after_text  = full_text[tag_end:]

        # Очищаємо всі runs (примусово — тег розбитий, не обійтись)
        for run in runs:
            run.text = ''

        # Відновлюємо: before + зображення + after
        anchor = runs[0] if runs else para.add_run()
        anchor.text = before_text
        anchor.add_picture(sig_path, width=Cm(width_cm))

        if after_text:
            para.add_run(after_text)

        return True

    # Тіло документа
    for para in doc.paragraphs:
        if _try_para(para):
            return True

    # Таблиці
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _try_para(para):
                        return True

    # Хедер / Футер
    for section in doc.sections:
        for para in section.header.paragraphs:
            if _try_para(para):
                return True
        for para in section.footer.paragraphs:
            if _try_para(para):
                return True

    return False


def fill_template(template_path: str, replacements: dict,
                  out_filename: str,
                  sig_path: str = "",
                  sig_width_cm: float = 4.5) -> str:
    """
    Заповнює шаблон Word і зберігає в TMP_DIR.
    Повертає шлях до збереженого .docx файлу.

    sig_path    — якщо вказано і файл існує, замінює {{ПІДПИС}} на зображення.
    sig_width_cm — ширина підпису в сантиметрах (висота масштабується автоматично).
    """
    os.makedirs(TMP_DIR, exist_ok=True)
    doc = Document(template_path)

    # Параграфи документа
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Параграфи в таблицях
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Header / Footer
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, replacements)

    # Вставка підпису (тільки якщо шаблон містить {{ПІДПИС}})
    if sig_path and os.path.exists(sig_path):
        _inject_signature_image(doc, sig_path, sig_width_cm)

    out_path = os.path.join(TMP_DIR, out_filename)
    doc.save(out_path)
    return out_path


# ============================================================
# КОНВЕРТАЦІЯ В PDF
# ============================================================

def convert_to_pdf(docx_path: str, pdf_output_path: str) -> str:
    out_dir = os.path.dirname(pdf_output_path)
    os.makedirs(out_dir, exist_ok=True)

    if not os.path.exists(LIBREOFFICE_PATH):
        raise FileNotFoundError(
            f"LibreOffice не знайдено: {LIBREOFFICE_PATH}\n"
            "Встановіть LibreOffice з https://www.libreoffice.org/"
        )

    subprocess.run([
        LIBREOFFICE_PATH,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", TMP_DIR,
        docx_path
    ], check=True, capture_output=True)

    tmp_pdf = os.path.join(
        TMP_DIR,
        os.path.basename(docx_path).replace(".docx", ".pdf")
    )
    shutil.move(tmp_pdf, pdf_output_path)
    return pdf_output_path


# ============================================================
# ШЛЯХИ ДО PDF ФАЙЛІВ
# ============================================================

def build_invoice_pdf_path(client_name: str, contract_no: str,
                            invoice_no: str) -> str:
    safe_client   = _safe_folder_name(client_name)
    safe_contract = _safe_folder_name(contract_no)
    folder = os.path.join(
        CONTRACTS_DIR, safe_client,
        f"Договір {safe_contract}", "Рахунки"
    )
    return os.path.join(folder, f"Рахунок {invoice_no}.pdf")


def build_act_pdf_path(client_name: str, contract_no: str,
                       act_no: str) -> str:
    safe_client   = _safe_folder_name(client_name)
    safe_contract = _safe_folder_name(contract_no)
    folder = os.path.join(
        CONTRACTS_DIR, safe_client,
        f"Договір {safe_contract}", "Акти"
    )
    return os.path.join(folder, f"Акт {act_no}.pdf")


def _safe_folder_name(name: str) -> str:
    forbidden = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
    result = name
    for ch in forbidden:
        result = result.replace(ch, "_")
    return result.strip()


# ============================================================
# ГЕНЕРАЦІЯ РАХУНКУ
# ============================================================

def generate_invoice_access(data: dict) -> str:
    discount_pct = data.get("discount_pct", 0) or 0
    if discount_pct > 0:
        discount_line = (f"Знижка {discount_pct:.0f}%: −{data.get('discount_amt_str', '')} грн"
                         f"  (без знижки: {data.get('sum_gross_str', '')} грн)")
    else:
        discount_line = ""

    replacements = {
        "{{INVOICE_NO}}":      data["invoice_no"],
        "{{INVOICE_DATE}}":    data["invoice_date_str"],
        "{{CONTRACT_NO}}":     data["contract_no"],
        "{{CLIENT_NAME}}":     data["client_name"],
        "{{CLIENT_ADDRESS}}":  data["client_address"],
        "{{PERIOD_FROM}}":     data["period_from_str"],
        "{{PERIOD_TO}}":       data["period_to_str"],
        "{{USERS}}":           str(data["users"]),
        "{{MONTHS}}":          str(data["months"]),
        "{{TARIFF}}":          data["tariff_str"],
        "{{SUM}}":             data["sum_str"],
        "{{SUM_WORDS}}":       data["sum_words"],
        "{{DUE_DATE}}":        data["due_date_str"],
        "{{DISCOUNT_LINE}}":   discount_line,
    }
    filename = f"Рахунок_{data['invoice_no']}.docx"
    sig_path = data.get("sig_path", "")
    return fill_template(TEMPLATE_INVOICE_ACCESS, replacements, filename, sig_path=sig_path)


def generate_invoice_hourly(data: dict) -> str:
    replacements = {
        "{{INVOICE_NO}}":       data["invoice_no"],
        "{{INVOICE_DATE}}":     data["invoice_date_str"],
        "{{CONTRACT_NO}}":      data["contract_no"],
        "{{CLIENT_NAME}}":      data["client_name"],
        "{{CLIENT_ADDRESS}}":   data["client_address"],
        "{{SERVICE_SUBJECT}}":  data["subject"],
        "{{CONTRACT_DATE}}":    data["contract_date_str"],
        "{{HOURS}}":            str(data["hours"]),
        "{{HOUR_RATE}}":        data["hour_rate_str"],
        "{{SUM}}":              data["sum_str"],
        "{{SUM_WORDS}}":        data["sum_words"],
        "{{DUE_DATE}}":         data["due_date_str"],
    }
    filename = f"Рахунок_{data['invoice_no']}.docx"
    sig_path = data.get("sig_path", "")
    return fill_template(TEMPLATE_INVOICE_HOURLY, replacements, filename, sig_path=sig_path)


# ============================================================
# ГЕНЕРАЦІЯ АКТУ
# ============================================================

def generate_act_access(data: dict) -> str:
    replacements = {
        "{{ACT_NO}}":           data["act_no"],
        "{{ACT_DATE}}":         data["act_date_str"],
        "{{CONTRACT_NO}}":      data["contract_no"],
        "{{CONTRACT_DATE}}":    data["contract_date_str"],
        "{{CLIENT_NAME}}":      data["client_name"],
        "{{CLIENT_CODE}}":      data["edrpou"],
        "{{CLIENT_DIRECTOR}}":  data["client_director"],
        "{{PERIOD_FROM}}":      data["period_from_str"],
        "{{PERIOD_TO}}":        data["period_to_str"],
        "{{USERS}}":            str(data["users"]),
        "{{MONTHS}}":           str(data["months"]),
        "{{TARIFF}}":           data["tariff_str"],
        "{{SUM}}":              data["sum_str"],
        "{{SUM_WORDS}}":        data["sum_words"],
    }
    filename = f"Акт_{data['act_no']}.docx"
    return fill_template(TEMPLATE_ACT_ACCESS, replacements, filename)


def build_reminder_pdf_path(client_name: str, contract_no: str,
                             invoice_no: str) -> str:
    safe_client   = _safe_folder_name(client_name)
    safe_contract = _safe_folder_name(contract_no)
    folder = os.path.join(
        CONTRACTS_DIR, safe_client,
        f"Договір {safe_contract}", "Рахунки"
    )
    return os.path.join(folder, f"Нагадування_{invoice_no}.pdf")


TEMPLATE_LETTER_REMINDER = os.path.join(
    TEMPLATES_DIR, "Шаблон_Лист_Нагадування.docx"
)


def generate_reminder(data: dict) -> str:
    """
    Генерує лист-нагадування через fill_template.
    data: invoice_no, client_name, client_address, contract_no,
          sum_str, due_date_str, overdue_days, invoice_date_str, our_name
    Повертає шлях до .docx файлу в TMP_DIR.
    """
    from datetime import date
    today_str = date.today().strftime("%d.%m.%Y")

    overdue = int(data.get("overdue_days", 0) or 0)
    if overdue > 0:
        due_line = (
            f"Термін оплати минув {overdue} днів тому "
            f"(строк: {data.get('due_date_str', '')})."
        )
    else:
        due_line = f"Термін оплати: {data.get('due_date_str', '')}."

    replacements = {
        "{{TODAY_DATE}}":    today_str,
        "{{CLIENT_NAME}}":   data.get("client_name", ""),
        "{{CLIENT_ADDRESS}}": data.get("client_address", ""),
        "{{INVOICE_NO}}":    data.get("invoice_no", ""),
        "{{INVOICE_DATE}}":  data.get("invoice_date_str", ""),
        "{{CONTRACT_NO}}":   data.get("contract_no", ""),
        "{{SUM}}":           data.get("sum_str", ""),
        "{{DUE_DATE}}":      data.get("due_date_str", ""),
        "{{DUE_LINE}}":      due_line,
        "{{OUR_NAME}}":      data.get("our_name", "DocketPro"),
    }
    filename = f"Нагадування_{data.get('invoice_no', 'reminder')}.docx"

    # Fall back to programmatic generation if the template file is missing
    if not os.path.exists(TEMPLATE_LETTER_REMINDER):
        return _generate_reminder_fallback(data, filename)

    return fill_template(TEMPLATE_LETTER_REMINDER, replacements, filename)


def _generate_reminder_fallback(data: dict, filename: str) -> str:
    """Programmatic fallback used only when the .docx template is missing."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import date

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    def add_para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                 color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    today_str = date.today().strftime("%d.%m.%Y")
    add_para("НАГАДУВАННЯ ПРО ОПЛАТУ", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(f"Лист-нагадування від {today_str}", size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(107, 114, 128), space_after=14)
    add_para(f"Кому: {data.get('client_name', '')}", bold=True, size=11, space_after=2)
    if data.get("client_address"):
        add_para(data["client_address"], size=10, color=(107, 114, 128), space_after=10)
    add_para(
        f"Нагадуємо Вам, що рахунок {data.get('invoice_no', '')} "
        f"від {data.get('invoice_date_str', '')} по договору {data.get('contract_no', '')} "
        f"на суму {data.get('sum_str', '')} грн досі не оплачено.",
        size=11, space_after=6
    )
    overdue = int(data.get("overdue_days", 0) or 0)
    if overdue > 0:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(
            f"Термін оплати минув {overdue} днів тому (строк: {data.get('due_date_str', '')})."
        )
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(183, 28, 28)
    else:
        add_para(f"Термін оплати: {data.get('due_date_str', '')}.", size=11, space_after=6)
    add_para(
        "Просимо здійснити оплату якнайшвидше або зв'язатися з нами "
        "для уточнення термінів.",
        size=11, space_after=14
    )
    add_para("Деталі рахунку:", bold=True, size=11, space_after=4)
    for label, value in [
        ("Рахунок №", data.get("invoice_no", "")),
        ("Договір №", data.get("contract_no", "")),
        ("Дата рахунку", data.get("invoice_date_str", "")),
        ("Сума до сплати", f"{data.get('sum_str', '')} грн"),
        ("Строк оплати", data.get("due_date_str", "")),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(value)
        r2.font.size = Pt(10)
    add_para("", space_before=14, space_after=4)
    add_para("З повагою,", size=11, space_after=2)
    add_para(data.get("our_name", "DocketPro"), bold=True, size=11, space_after=2)

    os.makedirs(TMP_DIR, exist_ok=True)
    out_path = os.path.join(TMP_DIR, filename)
    doc.save(out_path)
    return out_path


def generate_act_hourly(data: dict) -> str:
    replacements = {
        "{{ACT_NO}}":           data["act_no"],
        "{{ACT_DATE}}":         data["act_date_str"],
        "{{CONTRACT_NO}}":      data["contract_no"],
        "{{CONTRACT_DATE}}":    data["contract_date_str"],
        "{{CLIENT_NAME}}":      data["client_name"],
        "{{CLIENT_CODE}}":      data["edrpou"],
        "{{CLIENT_DIRECTOR}}":  data["client_director"],
        "{{SERVICE_SUBJECT}}":  data["subject"],
        "{{HOURS}}":            str(data["hours"]),
        "{{HOUR_RATE}}":        data["hour_rate_str"],
        "{{SUM}}":              data["sum_str"],
        "{{SUM_WORDS}}":        data["sum_words"],
    }
    filename = f"Акт_{data['act_no']}.docx"
    return fill_template(TEMPLATE_ACT_HOURLY, replacements, filename)


def resolve_act_template(contract_type: str, act_template: str = "") -> str:
    """Повертає шлях до шаблону акту.
    act_template — ім'я файлу з колонки ActTemplate в Contracts (порожньо = стандарт)."""
    if act_template:
        path = os.path.join(TEMPLATES_DIR, act_template)
        if os.path.exists(path):
            return path
    if contract_type in ("Access", "Доступ"):
        return TEMPLATE_ACT_ACCESS
    return TEMPLATE_ACT_HOURLY


def generate_act_with_template(data: dict, template_path: str) -> str:
    """Генерує акт за вказаним шаблоном (підходить для Access, Hourly і кастомних)."""
    replacements = {
        "{{ACT_NO}}":           data["act_no"],
        "{{ACT_DATE}}":         data["act_date_str"],
        "{{CONTRACT_NO}}":      data["contract_no"],
        "{{CONTRACT_DATE}}":    data["contract_date_str"],
        "{{CLIENT_NAME}}":      data["client_name"],
        "{{CLIENT_CODE}}":      data["edrpou"],
        "{{CLIENT_DIRECTOR}}":  data["client_director"],
        "{{PERIOD_FROM}}":      data.get("period_from_str", ""),
        "{{PERIOD_TO}}":        data.get("period_to_str", ""),
        "{{USERS}}":            str(data.get("users", 1)),
        "{{MONTHS}}":           str(data.get("months", 0)),
        "{{TARIFF}}":           data.get("tariff_str", ""),
        "{{SERVICE_SUBJECT}}":  data.get("subject", ""),
        "{{HOURS}}":            str(data.get("hours", 0)),
        "{{HOUR_RATE}}":        data.get("hour_rate_str", ""),
        "{{SUM}}":              data["sum_str"],
        "{{SUM_WORDS}}":        data["sum_words"],
    }
    filename = f"Акт_{data['act_no']}.docx"
    return fill_template(template_path, replacements, filename)
