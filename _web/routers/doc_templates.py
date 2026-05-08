"""Doc-Templates router — Word document templates management."""

import re
import shutil
from pathlib import Path

from fastapi import APIRouter, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from fastapi.templating import Jinja2Templates

from models import fmt_money, fmt_date
from database import get_db

router = APIRouter(prefix="/templates")
_jinja = Jinja2Templates(directory=str(Path(__file__).parent.parent / "templates"))
_jinja.env.globals["fmt_money"] = fmt_money
_jinja.env.globals["fmt_date"]  = fmt_date

# Where .docx template files live
TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "_templates"

# Standard template filenames (exact match → not custom)
STANDARD_NAMES = {
    "Шаблон_Рахунок_Доступ.docx",
    "Шаблон_Рахунок_Погодинно.docx",
    "Шаблон_Акт_Доступ.docx",
    "Шаблон_Акт_Погодинно.docx",
    "Шаблон_Лист_Нагадування.docx",
}

# ── Variable schemas ───────────────────────────────────────────

VARS_INVOICE_ACCESS = [
    "INVOICE_NO", "INVOICE_DATE", "CONTRACT_NO",
    "CLIENT_NAME", "CLIENT_ADDRESS",
    "PERIOD_FROM", "PERIOD_TO",
    "USERS", "MONTHS", "TARIFF",
    "SUM", "SUM_WORDS", "DUE_DATE", "DISCOUNT_LINE",
]
VARS_INVOICE_HOURLY = [
    "INVOICE_NO", "INVOICE_DATE", "CONTRACT_NO",
    "CLIENT_NAME", "CLIENT_ADDRESS",
    "SERVICE_SUBJECT", "CONTRACT_DATE",
    "HOURS", "HOUR_RATE",
    "SUM", "SUM_WORDS", "DUE_DATE",
]
VARS_ACT_ACCESS = [
    "ACT_NO", "ACT_DATE", "CONTRACT_NO", "CONTRACT_DATE",
    "CLIENT_NAME", "CLIENT_CODE", "CLIENT_DIRECTOR",
    "PERIOD_FROM", "PERIOD_TO",
    "USERS", "MONTHS", "TARIFF",
    "SUM", "SUM_WORDS",
]
VARS_ACT_HOURLY = [
    "ACT_NO", "ACT_DATE", "CONTRACT_NO", "CONTRACT_DATE",
    "CLIENT_NAME", "CLIENT_CODE", "CLIENT_DIRECTOR",
    "SERVICE_SUBJECT", "HOURS", "HOUR_RATE",
    "SUM", "SUM_WORDS",
]
VARS_LETTER = [
    "CLIENT_NAME", "CLIENT_ADDRESS",
    "INVOICE_NO", "INVOICE_DATE", "CONTRACT_NO",
    "SUM", "DUE_DATE", "DUE_LINE",
    "TODAY_DATE", "OUR_NAME",
]


def _classify(filename: str) -> dict:
    """Return metadata dict for a template file."""
    fn = filename
    name_lower = fn.lower()

    is_invoice = "рахунок" in name_lower
    is_act     = "акт" in name_lower
    is_access  = "доступ" in name_lower
    is_hourly  = "погодинно" in name_lower
    is_letter  = any(w in name_lower for w in ("лист", "нагадування", "letter", "reminder"))

    if is_invoice and is_access:
        category, category_label = "invoice", "Рахунок"
        ttype, ttype_label       = "access",  "Доступ"
        variables                = VARS_INVOICE_ACCESS
    elif is_invoice and is_hourly:
        category, category_label = "invoice", "Рахунок"
        ttype, ttype_label       = "hourly",  "Погодинно"
        variables                = VARS_INVOICE_HOURLY
    elif is_act and is_access:
        category, category_label = "act", "Акт"
        ttype, ttype_label       = "access", "Доступ"
        variables                = VARS_ACT_ACCESS
    elif is_act and is_hourly:
        category, category_label = "act", "Акт"
        ttype, ttype_label       = "hourly", "Погодинно"
        variables                = VARS_ACT_HOURLY
    elif is_letter:
        category, category_label = "letter", "Лист"
        ttype, ttype_label       = "letter", ""
        variables                = VARS_LETTER
    else:
        category, category_label = "other", "Інше"
        ttype, ttype_label       = "unknown", ""
        variables                = []

    # Human-readable display name
    display = fn
    for prefix in ("Шаблон_Рахунок_", "Шаблон_Акт_", "Шаблон_Лист_", "Шаблон_"):
        if display.startswith(prefix):
            display = display[len(prefix):]
            break
    if display.endswith(".docx"):
        display = display[:-5]
    display = display.replace("_", " ")

    size_kb = 0
    fpath = TEMPLATES_DIR / fn
    if fpath.exists():
        size_kb = round(fpath.stat().st_size / 1024)

    return {
        "filename":       fn,
        "display":        display,
        "category":       category,
        "category_label": category_label,
        "ttype":          ttype,
        "ttype_label":    ttype_label,
        "is_standard":    fn in STANDARD_NAMES,
        "variables":      variables,
        "var_count":      len(variables),
        "size_kb":        size_kb,
    }


def _load_templates() -> list[dict]:
    if not TEMPLATES_DIR.exists():
        return []
    files = sorted(TEMPLATES_DIR.glob("*.docx"))
    return [_classify(f.name) for f in files]


# ── Routes ────────────────────────────────────────────────────

@router.get("", response_class=HTMLResponse)
async def templates_page(request: Request):
    tmpl_list = _load_templates()
    counts = {
        "invoice": sum(1 for t in tmpl_list if t["category"] == "invoice"),
        "act":     sum(1 for t in tmpl_list if t["category"] == "act"),
        "letter":  sum(1 for t in tmpl_list if t["category"] == "letter"),
        "other":   sum(1 for t in tmpl_list if t["category"] == "other"),
    }

    # Load contracts for assignment widget
    with get_db() as db:
        contracts_rows = db.execute(
            "SELECT contract_no, client_name, contract_type, act_template "
            "FROM contracts WHERE status='Активний' ORDER BY client_name, contract_no"
        ).fetchall()

    all_contracts = [dict(r) for r in contracts_rows]

    # Build assignments map: filename → list of contracts using it
    assignments: dict[str, list[dict]] = {}
    for c in all_contracts:
        tpl = (c.get("act_template") or "").strip()
        if tpl:
            assignments.setdefault(tpl, []).append({
                "contract_no":   c["contract_no"],
                "client_name":   c["client_name"] or "",
                "contract_type": c["contract_type"] or "",
            })

    return _jinja.TemplateResponse("templates.html", {
        "request":      request,
        "templates":    tmpl_list,
        "total":        len(tmpl_list),
        "counts":       counts,
        "all_contracts": all_contracts,
        "assignments":   assignments,
    })


@router.get("/download/{filename}")
async def download_template(filename: str):
    filepath = TEMPLATES_DIR / filename
    if not filepath.exists() or not filepath.is_file():
        from fastapi.responses import JSONResponse
        return JSONResponse({"error": "not found"}, status_code=404)
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


@router.post("/upload", response_class=HTMLResponse)
async def upload_template(
    file: UploadFile = File(...),
    category: str    = Form("other"),
):
    """Upload a .docx template file to the _templates directory."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    filename = file.filename or "template.docx"
    # Sanitise filename — allow only safe characters (including Cyrillic)
    safe_name = "".join(
        c for c in filename
        if c.isalnum() or c in "._- "
    ).strip() or "template.docx"
    if not safe_name.endswith(".docx"):
        safe_name += ".docx"

    # Auto-prefix based on category if the keyword is not already in the filename
    name_lower = safe_name.lower()
    if category == "invoice" and "рахунок" not in name_lower:
        safe_name = "Шаблон_Рахунок_" + safe_name
    elif category == "act" and "акт" not in name_lower:
        safe_name = "Шаблон_Акт_" + safe_name
    elif category == "letter" and not any(
        w in name_lower for w in ("лист", "нагадування", "letter", "reminder")
    ):
        safe_name = "Шаблон_Лист_" + safe_name

    dest = TEMPLATES_DIR / safe_name
    with dest.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    return RedirectResponse("/templates", status_code=303)


# ── Helper ────────────────────────────────────────────────────

def _text_to_docx(text: str, path: Path) -> None:
    """Convert plain text (with optional # heading markers) to a .docx file."""
    from docx import Document  # type: ignore
    doc = Document()
    for line in text.split("\n"):
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


# ── New routes ────────────────────────────────────────────────

@router.post("/{filename}/delete", response_class=HTMLResponse)
async def delete_template(filename: str):
    """Delete a custom template file. Standard templates are protected."""
    if filename in STANDARD_NAMES:
        return JSONResponse({"error": "Cannot delete a standard template."}, status_code=403)

    # Path traversal protection
    dest = (TEMPLATES_DIR / filename).resolve()
    if not str(dest).startswith(str(TEMPLATES_DIR.resolve())):
        return JSONResponse({"error": "Invalid path."}, status_code=400)

    if dest.exists() and dest.is_file():
        dest.unlink()

    return RedirectResponse("/templates", status_code=303)


@router.get("/create", response_class=HTMLResponse)
async def create_template_page(request: Request):
    """Render the in-browser template editor (create mode)."""
    return _jinja.TemplateResponse("template_editor.html", {
        "request":         request,
        "editing":         False,
        "filename":        "",
        "initial_content": "",
    })


DOC_TYPE_PREFIXES = {
    "Рахунок Доступ":    "Шаблон_Рахунок_Доступ_",
    "Рахунок Погодинно": "Шаблон_Рахунок_Погодинно_",
    "Акт Доступ":        "Шаблон_Акт_Доступ_",
    "Акт Погодинно":     "Шаблон_Акт_Погодинно_",
    "Лист":              "Шаблон_Лист_",
}


@router.post("/create-save", response_class=HTMLResponse)
async def create_template_save(
    tpl_name: str = Form(...),
    doc_type: str = Form("Інший"),
    content:  str = Form(""),
):
    """Save the editor content as a new .docx template."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    # Sanitise the user-supplied name
    safe_tpl_name = re.sub(r"[^\wЀ-ӿԀ-ԯ\- ]", "", tpl_name).strip()
    safe_tpl_name = safe_tpl_name.replace(" ", "_") or "Шаблон"

    prefix = DOC_TYPE_PREFIXES.get(doc_type, "Шаблон_")
    filename = f"{prefix}{safe_tpl_name}.docx"

    dest = TEMPLATES_DIR / filename
    _text_to_docx(content, dest)

    return RedirectResponse("/templates", status_code=303)


@router.post("/{filename}/assign", response_class=HTMLResponse)
async def assign_template(
    filename:    str,
    contract_no: str = Form(...),
):
    """Assign this template to a contract (sets contracts.act_template)."""
    with get_db() as db:
        db.execute(
            "UPDATE contracts SET act_template=? WHERE contract_no=?",
            (filename, contract_no),
        )
    return RedirectResponse("/templates", status_code=303)


@router.post("/{filename}/unassign", response_class=HTMLResponse)
async def unassign_template(
    filename:    str,
    contract_no: str = Form(...),
):
    """Remove this template assignment from a contract."""
    with get_db() as db:
        db.execute(
            "UPDATE contracts SET act_template=NULL WHERE contract_no=? AND act_template=?",
            (contract_no, filename),
        )
    return RedirectResponse("/templates", status_code=303)
